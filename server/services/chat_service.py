import time
import numpy as np
from typing import Optional, Dict, List
from fastapi import HTTPException, UploadFile
from docx import Document
import pdfplumber
from dotenv import load_dotenv
from groq import Groq
import tempfile
import requests
import os
import ast
from PIL import Image
import pytesseract
from pdf2image import convert_from_path

import models.chat_model as chat_model

load_dotenv()

client_gradio = None
MAX_RETRIES = 10

def load_gradio_client():
    global client_gradio
    retries = 0
    from gradio_client import Client as GradioClient
    while retries < MAX_RETRIES and client_gradio is None:
        try:
            client_gradio = GradioClient("BienKieu/sentence-embedding")

            client_gradio.predict("Test")
        except Exception as e:
            print(f"Gradio client load failed ({retries+1}/{MAX_RETRIES}):", e)
            client_gradio = None
            retries += 1
            time.sleep(1)
    if client_gradio is None:
        print("Failed to load Gradio client after retries")

load_gradio_client()

def get_embeddings(texts):
    if client_gradio is None:
        raise RuntimeError("Gradio client is not available")
    
    res = client_gradio.predict(texts)  
    return np.array(res)  


def get_chat_sessions_by_user_id(user_id: int) -> list:
    chats = chat_model.get_chat_sessions_by_user_id(user_id)
    return chats if chats else []


def get_chat_session_by_id(chat_id: int) -> Dict:
    chat = chat_model.get_chat_session_by_id(chat_id)
    if not chat:
        raise HTTPException(status_code=404, detail="Chat session not found")
    return chat


def get_messages_by_chat_session_id(session_id: int) -> list:
    messages = chat_model.get_messages_by_session_id(session_id)
    if not messages:
        raise HTTPException(status_code=404, detail="No messages found for this chat session")
    return messages


def get_file_info_by_id(file_id: int) -> Dict:
    file = chat_model.get_file_by_id(file_id)
    if not file:
        raise HTTPException(status_code=404, detail="File not found")
    return file


def get_files_info_by_session_id(session_id: int) -> list:
    files = chat_model.get_files_by_session_id(session_id)
    if not files:
        raise HTTPException(status_code=404, detail="No files found for this chat session")
    return files


def get_chat_session_not_have_file(user_id: int) -> Dict:
    session = chat_model.get_chat_session_not_have_file(user_id)
    if not session:
        raise HTTPException(status_code=404, detail="No chat session without file found for this user")
    return session


def create_new_chat_session(user_id: int, title: str = "New Chat") -> Dict:
    if chat_model.count_chat_session_not_have_file(user_id) > 0:
        return chat_model.get_chat_session_not_have_file(user_id)
    return chat_model.create_chat_session(user_id, title)


def create_new_chat_message(session_id: int, role: str, content: str) -> Dict:
    return chat_model.create_chat_message(session_id, role, content)


def update_chat_session(session_id: int, new_name: str) -> Dict:
    if new_name.strip() == "":
        raise HTTPException(status_code=400, detail="Session name cannot be empty")
    result = chat_model.update_chat_session_title(session_id, new_name)
    if not result:
        raise HTTPException(status_code=404, detail="Session not found")
    return result

def read_pdf(path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                t = page.extract_text()
                if t and t.strip():
                    text += t + "\n"
                else:
                    try:
                        images = convert_from_path(path, first_page=page_num, last_page=page_num)
                        for img in images:
                            ocr_text = pytesseract.image_to_string(img, lang='eng')
                            if ocr_text.strip():
                                text += ocr_text + "\n"
                    except:
                        pass
    except:
        pass
    return text.strip()

def read_word(path: str) -> str:
    text_parts = []
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
    except:
        pass
    return "\n".join(text_parts).strip()

def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def read_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def chunk_text(text: str, chunk_size: int = 200, overlap: int = 20) -> List[str]:
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunks.append(" ".join(words[start:end]))
        start += chunk_size - overlap
    return chunks

def create_faiss_index_for_file(session_id: int, file_id: int, embedding_model_name: str = "all-MiniLM-L6-v2") -> bool:
    try:
        file_info = chat_model.get_file_by_id(file_id)
        if not file_info:
            return False

        file_url = file_info.get("file_url")
        if not file_url:
            return False

        r = requests.get(file_url)
        if r.status_code != 200:
            return False

        suffix = os.path.splitext(file_url)[-1].lower().split("?")[0]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(r.content)
            tmp_path = tmp.name

        text = ""
        if suffix == ".pdf":
            text = read_pdf(tmp_path)
        elif suffix == ".docx":
            text = read_word(tmp_path)
        elif suffix == ".txt":
            text = read_txt(tmp_path)
        elif suffix == ".md":
            text = read_md(tmp_path)
        else:
            os.unlink(tmp_path)
            return False

        os.unlink(tmp_path)
        if not text.strip():
            return False
        
        file_name = file_info.get("filename")

        chunks = chunk_text(text, chunk_size=200, overlap=20)
        embeddings = get_embeddings(chunks)
        if len(chunks) != len(embeddings):
            return False

        chunk_records = []
        for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            if emb is None:
                continue
            chunk_records.append({
                "session_id": session_id,
                "file_id": file_id,
                "file_name": file_name,
                "chunk_index": idx,
                "text": chunk,
                "embedding": emb.tolist(),
                "embedding_model": embedding_model_name,
                "chunk_size": len(chunk),
            })

        if not chunk_records:
            return False

        return chat_model.create_file_chunks(file_id, chunk_records)

    except Exception as e:
        print(f"Error in create_faiss_index_for_file: {e}")
        return False

def search_similar_chunks(session_id: int, query_embedding: np.ndarray, top_k: int = 5) -> List[Dict]:
    chunks = chat_model.get_file_chunks_by_session_id(session_id)
    if not chunks:
        return []

    valid_chunks = []
    embeddings_list = []

    for chunk in chunks:
        emb = chunk.get("embedding")
        if emb is None:
            continue
        if isinstance(emb, str):
            emb = np.array(ast.literal_eval(emb), dtype="float32")
        else:
            emb = np.array(emb, dtype="float32")
        if emb.shape[0] != len(query_embedding):
            continue
        embeddings_list.append(emb)
        valid_chunks.append(chunk)

    if not embeddings_list:
        return []

    embeddings = np.stack(embeddings_list)
    query_vec = np.array(query_embedding, dtype="float32")

    embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
    query_vec = query_vec / np.linalg.norm(query_vec)

    similarities = np.dot(embeddings, query_vec)
    top_indices = np.argsort(similarities)[::-1][:top_k]

    similar_chunks = []
    for idx in top_indices:
        chunk = valid_chunks[idx].copy()
        chunk["similarity_score"] = float(similarities[idx])
        similar_chunks.append(chunk)

    return similar_chunks


def get_query_embedding(query: str) -> np.ndarray:
    return get_embeddings([query])[0]


def call_llm_api(prompt) -> str:
    client_groq = Groq()
    completion = client_groq.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        top_p=0.9,
        stream=False,
        max_tokens=1000,
    )
    return completion.choices[0].message.content.strip()


def process_user_message(session_id: int, message: str) -> Dict:
    query_embedding = get_query_embedding(message)
    similar_chunks = search_similar_chunks(session_id, query_embedding, top_k=3)

    if similar_chunks:
        context = "\n\n".join([
            f"[From {chunk.get('file_name', 'unknown')}]: {chunk['text']}"
            for chunk in similar_chunks
        ])
    else:
        context = "No relevant documents were found in this session."

    prompt = f"""
User asks:
{message}

From the following documents (with file names shown):
{context}

Instructions:
- Mention the file name when citing information.
- Give a short but complete answer (concise, clear, not too long).
- Reply in the user's language.
- Only use the provided documents; do not add outside info.
- If nothing is relevant, reply: "No information found in the documents."
"""

    answer = call_llm_api(prompt)
    chat_model.create_chat_message(session_id, "bot", answer)

    return {"session_id": session_id, "user_message": message, "answer": answer}


def delete_chat_session_by_id(session_id: int) -> bool:
    return chat_model.delete_chat_session(session_id)


def delete_file_from_session(file_id: int) -> bool:
    return chat_model.delete_file_from_session(file_id)


def upload_file(user_id: int, session_id: int, uploaded_file: UploadFile) -> Dict:
    MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024  # 20 MB
    content = uploaded_file.file.read()
    if content is None:
        raise HTTPException(status_code=400, detail="Empty file upload")
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=413, detail="File too large. Maximum allowed size is 20 MB")

    class _PatchedUploadFile:
        def __init__(self, original: UploadFile, content_bytes: bytes):
            self.filename = original.filename
            self.file = type('F', (), {'read': lambda self2: content_bytes})()
            self.content_type = original.content_type

    patched_upload = _PatchedUploadFile(uploaded_file, content)
    res = chat_model.upload_file(user_id, patched_upload)
    if not res or "id" not in res:
        raise HTTPException(status_code=500, detail="File upload failed")

    file_id = res["id"]

    linked = chat_model.link_file_to_session(session_id, file_id)
    if not linked:
        raise HTTPException(status_code=500, detail="Failed to link file to session")

    faiss_created = create_faiss_index_for_file(session_id, file_id)
    if not faiss_created:
        raise HTTPException(status_code=500, detail="Failed to create FAISS index / file chunks")

    return res